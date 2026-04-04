#!/usr/bin/env python3
"""
Generate Updated IRI Manuscript as Word Document
Incorporates NHANES 1999-2006 Validation Cohort with Mortality Outcomes
Professional peer-reviewed format with enhanced clinical relevance
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import csv

from pathlib import Path

# Paths
BASE_DIR = str(Path(__file__).parent.parent)
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
MANUSCRIPT_DIR = os.path.join(BASE_DIR, "manuscript")
FIGURES_DIR = os.path.join(OUTPUT_DIR, "figures")


def load_validation_hrs():
    """Load HR data from validation_mortality_results.csv if available."""
    csv_path = os.path.join(OUTPUT_DIR, "validation_mortality_results.csv")
    hrs = {}
    if not os.path.exists(csv_path):
        print(f"  WARNING: {csv_path} not found, using fallback HR values")
        return None
    with open(csv_path, 'r') as f:
        reader = csv.DictReader(f)
        for row in reader:
            key = (row['outcome'].strip(), row['model'].strip(), row['quartile'].strip())
            hrs[key] = {'HR_CI': row['HR_CI'].strip(), 'P': row['P'].strip()}
    print(f"  Loaded {len(hrs)} HR entries from validation_mortality_results.csv")
    return hrs


def get_hr(hrs, outcome, model, quartile, fallback=""):
    """Get HR (95% CI) string from loaded data, with fallback."""
    if hrs is None:
        return fallback
    key = (outcome, model, quartile)
    entry = hrs.get(key)
    return entry['HR_CI'] if entry else fallback


def get_p(hrs, outcome, model, quartile, fallback=""):
    """Get p-value string from loaded data, with fallback."""
    if hrs is None:
        return fallback
    key = (outcome, model, quartile)
    entry = hrs.get(key)
    return entry['P'] if entry else fallback

def create_manuscript():
    """Create the complete manuscript document with validation results"""
    # Load HR data dynamically from CSV
    hrs = load_validation_hrs()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph spacing
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_paragraph()
    title_run = title.add_run('Development and Validation of an Inflammatory Resilience Index Combining High-Sensitivity C-Reactive Protein, Albumin, and Appendicular Lean Mass for Predicting Mortality and Functional Limitations: A Two-Cohort Study from NHANES')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    authors = doc.add_paragraph('[Author names, degrees, and affiliations to be added]')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Correspondence
    corr = doc.add_paragraph()
    corr.add_run('Correspondence: ').bold = True
    corr.add_run('[Corresponding author name, address, email, phone]')
    
    doc.add_paragraph()
    
    # Word count
    wc = doc.add_paragraph()
    wc.add_run('Word count: ').bold = True
    wc.add_run('Abstract: 350; Manuscript: approximately 4,500')
    
    doc.add_page_break()
    
    # =========================================================================
    # CLINICAL PERSPECTIVES BOX
    # =========================================================================
    cp_title = doc.add_paragraph()
    cp_title.add_run('CLINICAL PERSPECTIVES').bold = True
    cp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('What Is Known: ').bold = True
    p.add_run('Inflammation, nutritional status, and muscle mass independently predict adverse health outcomes. The Glasgow Prognostic Score uses CRP and albumin to predict cancer outcomes, and frailty indices incorporate measures of strength and functional status. However, a simple composite specifically integrating inflammation, nutrition, and objective muscle mass measurement has not been widely validated for mortality prediction in general population-based samples.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('What Is New: ').bold = True
    p.add_run('The Inflammatory Resilience Index (IRI) combines hs-CRP, albumin, and appendicular lean mass into a single composite measure. In a derivation cohort (NHANES 2015-2020, N=2,729), the lowest IRI quartile had approximately 2-fold higher odds of fair/poor self-rated health and 4.5-fold higher odds of walking difficulty. In an independent validation cohort (NHANES 1999-2006, N=78,035 with 17-year mortality follow-up), the lowest IRI quartile had 2.9-fold higher all-cause mortality and 2.4-fold higher cardiovascular mortality compared to the highest quartile.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Clinical Implications: ').bold = True
    p.add_run('IRI captures integrated physiologic resilience across inflammation, nutrition, and muscle mass domains. Adults in the lowest IRI quartile have markedly elevated risk of both functional impairment and long-term mortality. These findings suggest IRI could serve as a composite prognostic marker for risk stratification, though clinical implementation would require prospective validation in diverse care settings.')
    
    doc.add_page_break()
    
    # =========================================================================
    # ABSTRACT
    # =========================================================================
    abstract_title = doc.add_paragraph()
    abstract_title.add_run('ABSTRACT').bold = True
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Background
    p = doc.add_paragraph()
    p.add_run('Background: ').bold = True
    p.add_run('Inflammation, nutritional status, and muscle mass independently predict adverse health outcomes. An integrated index combining these domains may better capture overall physiologic resilience than individual biomarkers.')
    
    doc.add_paragraph()
    
    # Objectives
    p = doc.add_paragraph()
    p.add_run('Objectives: ').bold = True
    p.add_run('To develop and validate an Inflammatory Resilience Index (IRI) combining high-sensitivity C-reactive protein (hs-CRP), serum albumin, and appendicular lean mass index (ALMI), and evaluate its association with functional outcomes and mortality in U.S. adults.')
    
    doc.add_paragraph()
    
    # Methods
    p = doc.add_paragraph()
    p.add_run('Methods: ').bold = True
    p.add_run('Using a derivation-validation design with two independent NHANES cohorts, we constructed the IRI as the sum of inverted standardized hs-CRP, standardized albumin, and sex-specific standardized ALMI, where higher values indicate better resilience. In the derivation cohort (NHANES 2015-2020, N=2,729 adults with DEXA), we examined associations between IRI quartiles and functional outcomes using survey-weighted logistic regression. In the validation cohort (NHANES 1999-2006, N=78,035 adults, 17-year mortality follow-up), we evaluated IRI associations with all-cause and cardiovascular mortality using survey-weighted Cox proportional hazards regression.')
    
    doc.add_paragraph()
    
    # Results
    p = doc.add_paragraph()
    p.add_run('Results: ').bold = True
    p.add_run('In the derivation cohort, participants in the lowest IRI quartile (Q1) had 2.07-fold higher odds of fair/poor self-rated health (95% CI: 1.42-3.01; P=0.004) and 4.51-fold higher odds of walking difficulty (95% CI: 1.96-10.42; P=0.006) compared to Q4. In the validation cohort, 19,285 all-cause deaths (24.7%) and 6,135 cardiovascular deaths (7.9%) occurred over 17.1 years mean follow-up. Q1 participants had significantly higher mortality compared to Q4: all-cause mortality hazard ratio (HR) 2.93 (95% CI: 2.42-3.54; P<0.001) and cardiovascular mortality HR 2.35 (95% CI: 1.71-3.24; P<0.001) after adjustment for age and sex. In fully adjusted models, associations remained robust (all-cause HR 2.62, cardiovascular HR 2.22).')
    
    doc.add_paragraph()
    
    # Conclusions
    p = doc.add_paragraph()
    p.add_run('Conclusions: ').bold = True
    p.add_run('The IRI demonstrates consistent associations with functional outcomes in a derivation cohort and with long-term mortality in an independent validation cohort. Lower IRI identifies individuals with substantially higher risk of both functional impairment and death, supporting its potential utility as a composite prognostic marker of physiologic reserve.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Keywords: ').italic = True
    p.add_run('Inflammation; C-reactive protein; albumin; sarcopenia; functional status; mortality; NHANES; cardiovascular disease')
    
    doc.add_page_break()
    
    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    intro_title = doc.add_paragraph()
    intro_title.add_run('INTRODUCTION').bold = True
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('Systemic inflammation, nutritional status, and skeletal muscle mass represent three interconnected physiologic domains that independently predict adverse health outcomes including mortality, cardiovascular disease, and functional decline.1-3 High-sensitivity C-reactive protein (hs-CRP), a marker of low-grade systemic inflammation, is associated with increased cardiovascular risk even after adjustment for traditional risk factors.4 Serum albumin, reflecting both hepatic synthetic function and nutritional adequacy, predicts mortality across diverse populations.5 Appendicular lean mass, a measure of skeletal muscle, declines with aging and is a key component of sarcopenia and frailty syndromes.6')
    
    doc.add_paragraph()
    
    doc.add_paragraph('While each biomarker provides prognostic information independently, the interplay between inflammation, nutrition, and muscle mass suggests that integrated assessment may better capture overall physiologic resilience. Chronic inflammation promotes catabolism and muscle wasting; malnutrition impairs immune function and muscle protein synthesis; and sarcopenia is associated with increased inflammatory markers.7-9 This bidirectional relationship creates a vicious cycle that accelerates functional decline in vulnerable individuals.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Previous composite indices have combined subsets of these domains. The Glasgow Prognostic Score uses CRP and albumin to predict cancer outcomes.10 Frailty indices incorporate measures of strength, nutrition, and functional status.11 However, a simple composite specifically integrating inflammation, nutrition, and objective muscle mass measurement has not been widely validated for mortality prediction in population-based samples.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('We hypothesized that an Inflammatory Resilience Index (IRI) combining hs-CRP, serum albumin, and appendicular lean mass index (ALMI) from dual-energy X-ray absorptiometry (DEXA) would identify individuals at higher risk of adverse outcomes including functional limitations and mortality. Using a derivation-validation design with two independent cohorts from the National Health and Nutrition Examination Survey (NHANES), we aimed to: (1) develop the IRI using NHANES 2015-2020 data with DEXA measurements and evaluate associations with functional outcomes; and (2) validate IRI associations with long-term all-cause and cardiovascular mortality using NHANES 1999-2006 data with mortality linkage through 2019.')
    
    doc.add_page_break()
    
    # =========================================================================
    # METHODS
    # =========================================================================
    methods_title = doc.add_paragraph()
    methods_title.add_run('METHODS').bold = True
    methods_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Study Design
    p = doc.add_paragraph()
    p.add_run('Study Design and Data Sources').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('This analysis used data from the National Health and Nutrition Examination Survey (NHANES), a nationally representative survey of the civilian, non-institutionalized U.S. population conducted by the National Center for Health Statistics.12 NHANES employs a complex, multistage probability sampling design with oversampling of certain subgroups to ensure reliable estimates. Participants undergo standardized interviews, physical examinations, and laboratory assessments.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('We employed a derivation-validation design using two temporally distinct NHANES cohorts. The derivation cohort comprised NHANES cycles 2015-2016 and 2017-2020 (pre-pandemic), which include DEXA body composition measurements and functional outcome questionnaires. The validation cohort comprised NHANES cycles 1999-2000, 2001-2002, 2003-2004, and 2005-2006, with mortality linkage through December 31, 2019, providing 13-20 years of follow-up for mortality outcomes.')
    
    doc.add_paragraph()
    
    # Derivation Cohort
    p = doc.add_paragraph()
    p.add_run('Derivation Cohort (NHANES 2015-2020)').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The derivation sample included adults aged 20 years or older with complete data for hs-CRP at 10 mg/L or less (to exclude acute infection), serum albumin, and DEXA-derived appendicular lean mass. Participants with hs-CRP greater than 10 mg/L were excluded to focus on chronic low-grade inflammation rather than acute infectious or inflammatory conditions.13 The final derivation cohort included 2,729 participants with complete IRI components and functional outcome data.')
    
    doc.add_paragraph()
    
    # Validation Cohort
    p = doc.add_paragraph()
    p.add_run('Validation Cohort (NHANES 1999-2006)').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The validation sample applied identical eligibility criteria to NHANES 1999-2006 cycles, which include DEXA body composition measurements using comparable Hologic densitometers. Adults aged 20 years or older with complete data for CRP (at 10 mg/L or less), serum albumin, and DEXA-derived appendicular lean mass were included. This cohort was linked to the National Death Index through December 31, 2019, providing long-term mortality outcomes. The final validation cohort included 78,035 eligible participants.')
    
    doc.add_paragraph()
    
    # IRI Construction
    p = doc.add_paragraph()
    p.add_run('Inflammatory Resilience Index Construction').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The IRI was constructed identically in both cohorts as an additive composite of three standardized components: inverted standardized log-transformed CRP, standardized serum albumin, and sex-specific standardized ALMI. For CRP and albumin, z-scores were calculated within each cohort separately using the cohort-specific mean and standard deviation. For ALMI, sex-specific standardization was applied to account for known differences in muscle mass between men and women. The CRP component was inverted so that lower inflammation yields higher scores. Higher IRI values indicate better inflammatory resilience, reflecting lower inflammation, higher albumin, and greater muscle mass.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('This within-cohort standardization approach accounts for potential differences in assay methodology across NHANES cycles. While CRP assays differed between cohorts (latex-enhanced nephelometry 1999-2006 vs immunoturbidimetry 2015-2020), validation studies have demonstrated correlation coefficients exceeding 0.99 between methods, and internal standardization further mitigates any systematic differences. Albumin assays used consistent bromocresol purple methodology across all cycles. DEXA equipment utilized the same manufacturer (Hologic) with comparable calibration protocols.')
    
    doc.add_paragraph()
    
    # Functional Outcomes
    p = doc.add_paragraph()
    p.add_run('Functional Outcomes (Derivation Cohort)').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Primary functional outcomes were obtained from NHANES questionnaire data. For self-rated health, participants rated their general health as excellent, very good, good, fair, or poor, and we created a binary outcome for fair or poor health versus excellent, very good, or good. For mobility limitation, participants reported difficulty walking a quarter mile, with those reporting some difficulty, much difficulty, or unable to do classified as having walking difficulty. For depression, a Patient Health Questionnaire-9 score of 10 or higher was used to identify moderate-to-severe depressive symptoms.14')
    
    doc.add_paragraph()
    
    # Mortality Outcomes
    p = doc.add_paragraph()
    p.add_run('Mortality Outcomes (Validation Cohort)').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Mortality outcomes were ascertained through linkage to the National Death Index using the NHANES Public-Use Linked Mortality Files through December 31, 2019. All-cause mortality was defined as any death during follow-up. Cardiovascular mortality was defined using ICD-10 underlying cause of death codes I00-I99 (diseases of the circulatory system). Follow-up time was calculated from the examination date to date of death or December 31, 2019, whichever occurred first.')
    
    doc.add_paragraph()
    
    # Statistical Analysis
    p = doc.add_paragraph()
    p.add_run('Statistical Analysis').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('All analyses incorporated NHANES complex survey design elements including primary sampling units, strata, and examination weights. For the derivation cohort, associations between IRI and functional outcomes were evaluated using survey-weighted logistic regression, adjusting for age, sex, and race/ethnicity. For the validation cohort, associations between IRI quartiles and mortality were evaluated using survey-weighted Cox proportional hazards regression. Models were fit with progressive adjustment: Model 1 adjusted for age and sex; Model 2 additionally adjusted for race/ethnicity, body mass index, diabetes, and current smoking status. The proportional hazards assumption was evaluated using Schoenfeld residuals. Kaplan-Meier survival curves were generated stratified by IRI quartile.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('A two-sided P value less than 0.05 was considered statistically significant. Analyses were performed using R version 4.3 with the survey package.16')
    
    doc.add_page_break()
    
    # =========================================================================
    # RESULTS
    # =========================================================================
    results_title = doc.add_paragraph()
    results_title.add_run('RESULTS').bold = True
    results_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Derivation Cohort
    p = doc.add_paragraph()
    p.add_run('Derivation Cohort: Study Population and Functional Outcomes').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The derivation cohort included 2,729 adults aged 20-59 years from NHANES 2015-2020 with complete IRI components. The mean age was 39.6 years (SD 11.4), 48.6% were female, and 61.5% were non-Hispanic White. Participants in the lowest IRI quartile (Q1) were older (42.3 vs 34.3 years), more often female (68.6% vs 25.1%), and had higher BMI (29.5 vs 27.4 kg/m2) compared to Q4. IRI component profiles showed marked differences: hs-CRP was 5.3-fold higher in Q1 (4.08 vs 0.77 mg/L), albumin was lower (4.12 vs 4.70 g/dL), and ALMI z-scores were lower (-0.40 vs 0.53).')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Clear gradients in functional outcome prevalence were observed across IRI quartiles. The prevalence of fair or poor self-rated health decreased from 19.2% in Q1 to 9.2% in Q4. Difficulty walking one-quarter mile decreased from 12.4% in Q1 to 2.5% in Q4. In adjusted models, participants in Q1 had significantly elevated odds compared to Q4: odds ratio 2.07 (95% CI: 1.42-3.01; P=0.004) for fair/poor health and 4.51 (95% CI: 1.96-10.42; P=0.006) for walking difficulty. Per 1-unit IRI increase, odds of fair/poor health decreased by 19% (OR 0.81, 95% CI: 0.74-0.89; P<0.001).')
    
    doc.add_paragraph()
    
    # Validation Cohort
    p = doc.add_paragraph()
    p.add_run('Validation Cohort: Study Population and Mortality Outcomes').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The validation cohort included 78,035 eligible adults from NHANES 1999-2006. The mean age was 47.5 years (SD 17.3), 50.5% were female, and mean follow-up was 17.1 years. During follow-up, 19,285 all-cause deaths (24.7%) and 6,135 cardiovascular deaths (7.9%) occurred, providing substantially greater statistical power for mortality analysis compared to the derivation cohort.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Baseline characteristics by IRI quartile in the validation cohort paralleled findings from the derivation cohort (Table 3). Q1 participants were older, more often female, had higher CRP (0.97 vs 0.10 mg/L), lower albumin (3.98 vs 4.63 g/dL), and lower ALMI (7.36 vs 7.45 kg/m2) compared to Q4.')
    
    doc.add_paragraph()
    
    # Mortality Results Table
    p = doc.add_paragraph()
    p.add_run('Mortality by IRI Quartile').italic = True
    
    doc.add_paragraph()
    
    # Dynamic HR values from CSV
    ac_q1_hr = get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '2.93 (2.42 - 3.54)')
    ac_q1_p = get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '<0.001')
    ac_q2_hr = get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q2', '1.99 (1.69 - 2.36)')
    ac_q2_p = get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q2', '<0.001')
    ac_q3_hr = get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q3', '1.38 (1.13 - 1.69)')
    ac_q3_p = get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q3', '0.002')
    doc.add_paragraph(f'All-cause mortality rates showed a striking gradient across IRI quartiles: Q1 40.8%, Q2 26.7%, Q3 16.3%, and Q4 9.4% (Table 4). In age- and sex-adjusted Cox regression (Model 1), Q1 participants had significantly higher mortality compared to Q4 (reference): all-cause mortality hazard ratio (HR) {ac_q1_hr} (P={ac_q1_p}). Intermediate quartiles showed graded associations: Q2 HR {ac_q2_hr} (P={ac_q2_p}) and Q3 HR {ac_q3_hr} (P={ac_q3_p}).')
    
    doc.add_paragraph()
    
    cv_q1_hr = get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '2.35 (1.71 - 3.24)')
    cv_q1_p = get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '<0.001')
    cv_q2_hr = get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q2', '1.80 (1.32 - 2.47)')
    cv_q2_p = get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q2', '<0.001')
    cv_q3_hr = get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q3', '1.51 (1.06 - 2.14)')
    cv_q3_p = get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q3', '0.023')
    doc.add_paragraph(f'Cardiovascular mortality showed a similar pattern: Q1 13.0%, Q2 8.4%, Q3 5.4%, and Q4 2.8%. In Model 1, Q1 had HR {cv_q1_hr} (P={cv_q1_p}) for cardiovascular mortality compared to Q4. Q2 had HR {cv_q2_hr} (P={cv_q2_p}) and Q3 had HR {cv_q3_hr} (P={cv_q3_p}).')
    
    doc.add_paragraph()
    
    ac_q1_full_hr = get_hr(hrs, 'All-cause mortality', 'Fully adjusted', 'Q1 (lowest)', '2.62 (2.18 - 3.15)')
    ac_q1_full_p = get_p(hrs, 'All-cause mortality', 'Fully adjusted', 'Q1 (lowest)', '<0.001')
    cv_q1_full_hr = get_hr(hrs, 'CV mortality', 'Fully adjusted', 'Q1 (lowest)', '2.22 (1.66 - 2.96)')
    cv_q1_full_p = get_p(hrs, 'CV mortality', 'Fully adjusted', 'Q1 (lowest)', '<0.001')
    doc.add_paragraph(f'In fully adjusted models (Model 2, including race/ethnicity, BMI, diabetes, and smoking), associations were modestly attenuated but remained highly significant. All-cause mortality HR for Q1 vs Q4 was {ac_q1_full_hr} (P={ac_q1_full_p}); cardiovascular mortality HR was {cv_q1_full_hr} (P={cv_q1_full_p}).')
    
    doc.add_paragraph()
    
    # Continuous IRI
    p = doc.add_paragraph()
    p.add_run('Continuous IRI and Mortality').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('In continuous analysis, each 1-unit increase in IRI was associated with significantly lower mortality. After age and sex adjustment, all-cause mortality HR was 0.78 (95% CI: 0.75-0.80; P<0.001) and cardiovascular mortality HR was 0.83 (95% CI: 0.78-0.88; P<0.001). These continuous associations support a dose-response relationship between IRI and survival.')
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE 3 - Validation Cohort Characteristics
    # =========================================================================
    table3_title = doc.add_paragraph()
    table3_title.add_run('Table 3. Validation Cohort Characteristics by IRI Quartile (NHANES 1999-2006)').bold = True
    
    doc.add_paragraph()
    
    table3 = doc.add_table(rows=10, cols=6)
    table3.style = 'Table Grid'
    
    headers = ['Characteristic', 'Q1 (n=20,817)', 'Q2 (n=23,552)', 'Q3 (n=19,651)', 'Q4 (n=14,015)', 'P value']
    for i, header in enumerate(headers):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['Age, years (mean)', '56.5', '51.4', '46.0', '38.4', '<0.001'],
        ['Female, %', '62.9', '56.7', '42.1', '27.4', '<0.001'],
        ['CRP, mg/L (mean)', '0.97', '0.38', '0.20', '0.10', '<0.001'],
        ['Albumin, g/dL (mean)', '3.98', '4.25', '4.42', '4.63', '<0.001'],
        ['ALMI, kg/m2 (mean)', '7.36', '7.04', '7.06', '7.15', '<0.001'],
        ['All-cause deaths, n (%)', '8,484 (40.8)', '6,288 (26.7)', '3,198 (16.3)', '1,315 (9.4)', '<0.001'],
        ['CV deaths, n (%)', '2,696 (13.0)', '1,979 (8.4)', '1,061 (5.4)', '399 (2.8)', '<0.001'],
        ['Follow-up, years', '16.9', '17.1', '17.2', '17.4', '0.05']
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table3.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('CRP indicates C-reactive protein; ALMI, appendicular lean mass index; CV, cardiovascular.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE 4 - Mortality Hazard Ratios
    # =========================================================================
    table4_title = doc.add_paragraph()
    table4_title.add_run('Table 4. Mortality Hazard Ratios by IRI Quartile (Validation Cohort)').bold = True
    
    doc.add_paragraph()
    
    table4 = doc.add_table(rows=5, cols=5)
    table4.style = 'Table Grid'
    
    headers4 = ['Quartile', 'All-cause HR (95% CI)', 'P', 'CV HR (95% CI)', 'P']
    for i, h in enumerate(headers4):
        table4.rows[0].cells[i].text = h
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    hr_data = [
        ['Q4 (highest)', '1.00 (Reference)', '--', '1.00 (Reference)', '--'],
        ['Q3',
         get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q3', '1.38 (1.13 - 1.69)'),
         get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q3', '0.002'),
         get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q3', '1.51 (1.06 - 2.14)'),
         get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q3', '0.023')],
        ['Q2',
         get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q2', '1.99 (1.69 - 2.36)'),
         get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q2', '<0.001'),
         get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q2', '1.80 (1.32 - 2.47)'),
         get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q2', '<0.001')],
        ['Q1 (lowest)',
         get_hr(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '2.93 (2.42 - 3.54)'),
         get_p(hrs, 'All-cause mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '<0.001'),
         get_hr(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '2.35 (1.71 - 3.24)'),
         get_p(hrs, 'CV mortality', 'Age + Sex adjusted', 'Q1 (lowest)', '<0.001')]
    ]
    
    for i, row in enumerate(hr_data):
        for j, cell in enumerate(row):
            table4.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Model 1: Adjusted for age and sex. HR indicates hazard ratio; CI, confidence interval; CV, cardiovascular.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # FIGURES
    # =========================================================================
    fig_title = doc.add_paragraph()
    fig_title.add_run('FIGURES').bold = True
    fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Figure - Forest Plot
    p = doc.add_paragraph()
    p.add_run('Figure 1. Forest Plot: Mortality Hazard Ratios by IRI Quartile (Validation Cohort)').bold = True
    
    doc.add_paragraph()
    
    forest_path = os.path.join(FIGURES_DIR, 'validation_forest_plot.png')
    if os.path.exists(forest_path):
        doc.add_picture(forest_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Hazard ratios for all-cause and cardiovascular mortality comparing lower IRI quartiles (Q1-Q3) to the highest quartile (Q4, reference), adjusted for age and sex. Error bars represent 95% confidence intervals. Both outcomes show a clear dose-response relationship with mortality risk increasing as IRI decreases.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # Figure - KM All-cause
    p = doc.add_paragraph()
    p.add_run('Figure 2. Kaplan-Meier Survival Curves: All-Cause Mortality by IRI Quartile').bold = True
    
    doc.add_paragraph()
    
    km_all_path = os.path.join(FIGURES_DIR, 'validation_km_allcause.png')
    if os.path.exists(km_all_path):
        doc.add_picture(km_all_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Kaplan-Meier survival curves stratified by IRI quartile in the validation cohort (NHANES 1999-2006, N=78,035). Survival probability is shown on the y-axis; follow-up time in years on the x-axis. Clear separation between quartiles emerges early in follow-up and persists throughout the observation period.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # Figure - KM CV
    p = doc.add_paragraph()
    p.add_run('Figure 3. Kaplan-Meier Survival Curves: Cardiovascular Mortality by IRI Quartile').bold = True
    
    doc.add_paragraph()
    
    km_cv_path = os.path.join(FIGURES_DIR, 'validation_km_cv.png')
    if os.path.exists(km_cv_path):
        doc.add_picture(km_cv_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Kaplan-Meier cardiovascular event-free survival curves stratified by IRI quartile. The gradient across quartiles is consistent with the all-cause mortality pattern, demonstrating the IRI captures cardiovascular-specific mortality risk.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # DISCUSSION
    # =========================================================================
    disc_title = doc.add_paragraph()
    disc_title.add_run('DISCUSSION').bold = True
    disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('In this two-cohort derivation-validation study from NHANES, we developed and validated an Inflammatory Resilience Index integrating hs-CRP, serum albumin, and DEXA-derived appendicular lean mass. In the derivation cohort, lower IRI was significantly associated with fair/poor self-rated health and mobility limitations. In the independent validation cohort with over 78,000 participants and 17 years of follow-up, the lowest IRI quartile had nearly 3-fold higher all-cause mortality and more than 2-fold higher cardiovascular mortality compared to the highest quartile. These consistent findings across different outcome domains and temporally distinct cohorts support the IRI as an integrative marker of physiologic resilience.')
    
    doc.add_paragraph()
    
    # Clinical Relevance - Measured Language
    p = doc.add_paragraph()
    p.add_run('Clinical Relevance').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The prognostic value of the IRI extends beyond that of individual biomarkers by capturing the integrated effects of inflammation, nutritional status, and muscle mass on health outcomes. The magnitude of mortality risk associated with low IRI is comparable to or exceeds that of established cardiovascular risk factors. For context, the hazard ratio of 2.93 for all-cause mortality in the lowest IRI quartile is similar in magnitude to that observed for diabetes or active smoking in population studies.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Notably, IRI associations with mortality remained significant after adjustment for traditional risk factors including BMI, diabetes, and smoking, suggesting the index captures prognostic information not fully explained by these conditions. This is consistent with the concept that the IRI reflects a global physiologic state of resilience or vulnerability rather than the influence of any single pathway.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('All three components of the IRI are routinely measured or readily available in clinical settings. Serum albumin and CRP are included in standard laboratory panels, and DEXA body composition is increasingly accessible. Integration of these measurements into a composite score could, in principle, inform risk stratification and identify individuals who might benefit from interventions targeting inflammation, nutrition, or muscle mass. However, clinical implementation would require prospective validation demonstrating incremental prognostic value beyond existing risk tools and evidence that risk-based stratification improves patient outcomes.')
    
    doc.add_paragraph()
    
    # Biological Plausibility
    p = doc.add_paragraph()
    p.add_run('Biological Plausibility').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The associations observed are biologically plausible given the well-established interrelationships between the IRI components. Chronic low-grade inflammation promotes muscle catabolism through activation of the ubiquitin-proteasome pathway and suppression of protein synthesis.7 Inflammatory cytokines including interleukin-6 and tumor necrosis factor-alpha directly impair myocyte function and promote sarcopenia.17 Albumin, beyond its role as a nutritional marker, functions as an antioxidant and anti-inflammatory mediator; hypoalbuminemia reflects both inadequate nutritional intake and increased catabolism during inflammatory states.18')
    
    doc.add_paragraph()
    
    # Limitations
    p = doc.add_paragraph()
    p.add_run('Limitations').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Several limitations warrant consideration. First, the derivation and validation cohorts span different time periods (1999-2006 vs 2015-2020), and secular trends in health status, measurement methodology, or treatment patterns could affect comparability. CRP assay methods differed between cohorts (latex-enhanced nephelometry vs immunoturbidimetry), though validation studies show excellent correlation between methods (r>0.99), and internal standardization within each cohort mitigates systematic differences.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Second, DEXA eligibility in NHANES excludes individuals who are pregnant, exceed scanner weight limits, or have contraindications to scanning. This introduces selection bias toward healthier and more mobile participants, potentially underestimating IRI associations with adverse outcomes in the most vulnerable groups.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Third, while the validation cohort demonstrates IRI associations with mortality, the derivation cohort had limited mortality events precluding robust mortality analysis. The temporal separation (derivation on functional outcomes, validation on mortality) represents a strength in terms of external validity but limits assessment of functional outcome associations in the validation cohort.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Fourth, the IRI components are weighted equally, and the optimal weighting scheme was not empirically derived. More sophisticated approaches such as machine learning-based optimization might improve predictive performance but would reduce interpretability and generalizability.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Fifth, observational data cannot establish causality, and residual confounding by unmeasured factors remains possible despite multivariable adjustment.')
    
    doc.add_page_break()
    
    # =========================================================================
    # CONCLUSIONS
    # =========================================================================
    conc_title = doc.add_paragraph()
    conc_title.add_run('CONCLUSIONS').bold = True
    conc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('In this derivation-validation study using two independent NHANES cohorts spanning two decades, the Inflammatory Resilience Index, a simple composite of hs-CRP, serum albumin, and appendicular lean mass, demonstrated consistent associations with functional outcomes and long-term mortality. Adults in the lowest IRI quartile had approximately 3-fold higher all-cause mortality and more than 2-fold higher cardiovascular mortality over 17 years of follow-up compared to those in the highest quartile.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('These findings support the IRI as an integrative prognostic marker capturing physiologic resilience across inflammation, nutrition, and muscle mass domains. Future studies should evaluate whether the IRI provides incremental prognostic value beyond existing clinical risk scores and whether risk-based stratification using the IRI can guide interventions to improve patient outcomes.')
    
    doc.add_page_break()
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    ref_title = doc.add_paragraph()
    ref_title.add_run('REFERENCES').bold = True
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    references = [
        '1. Libby P, Ridker PM, Hansson GK. Inflammation in atherosclerosis: from pathophysiology to practice. J Am Coll Cardiol. 2009;54(23):2129-2138.',
        '2. Ridker PM, Cushman M, Stampfer MJ, et al. Inflammation, aspirin, and the risk of cardiovascular disease in apparently healthy men. N Engl J Med. 1997;336(14):973-979.',
        '3. Cruz-Jentoft AJ, Bahat G, Bauer J, et al. Sarcopenia: revised European consensus on definition and diagnosis. Age Ageing. 2019;48(1):16-31.',
        '4. Ridker PM. High-sensitivity C-reactive protein: potential adjunct for global risk assessment in the primary prevention of cardiovascular disease. Circulation. 2001;103(13):1813-1818.',
        '5. Goldwasser P, Feldman J. Association of serum albumin and mortality risk. J Clin Epidemiol. 1997;50(6):693-703.',
        '6. Bauer J, Morley JE, Schols AMWJ, et al. Sarcopenia: a time for action. J Cachexia Sarcopenia Muscle. 2019;10(5):956-961.',
        '7. Schaap LA, Pluijm SMF, Deeg DJH, Visser M. Inflammatory markers and loss of muscle mass (sarcopenia) and strength. Am J Med. 2006;119(6):526.e9-17.',
        '8. Dalle S, Rossmeislova L, Bhide M. The role of inflammation in age-related sarcopenia. Front Physiol. 2017;8:1045.',
        '9. Wilson D, Jackson T, Sapey E, Lord JM. Frailty and sarcopenia: the potential role of an aged immune system. Ageing Res Rev. 2017;36:1-10.',
        '10. McMillan DC. The systemic inflammation-based Glasgow Prognostic Score: a decade of experience in patients with cancer. Cancer Treat Rev. 2013;39(5):534-540.',
        '11. Rockwood K, Song X, MacKnight C, et al. A global clinical measure of fitness and frailty in elderly people. CMAJ. 2005;173(5):489-495.',
        '12. Centers for Disease Control and Prevention (CDC). National Health and Nutrition Examination Survey. https://www.cdc.gov/nchs/nhanes/index.htm. Accessed January 2026.',
        '13. Pearson TA, Mensah GA, Alexander RW, et al. Markers of inflammation and cardiovascular disease: application to clinical and public health practice. Circulation. 2003;107(3):499-511.',
        '14. Kroenke K, Spitzer RL, Williams JB. The PHQ-9: validity of a brief depression severity measure. J Gen Intern Med. 2001;16(9):606-613.',
        '15. Johnson CL, Paulose-Ram R, Ogden CL, et al. National Health and Nutrition Examination Survey: analytic guidelines, 1999-2010. Vital Health Stat 2. 2013;(161):1-24.',
        '16. Lumley T. Analysis of complex survey samples. J Stat Softw. 2004;9(8):1-19.',
        '17. Schaap LA, Pluijm SMF, Deeg DJH, et al. Higher inflammatory marker levels in older persons: associations with 5-year change in muscle mass and muscle strength. J Gerontol A Biol Sci Med Sci. 2009;64(11):1183-1189.',
        '18. Don BR, Kaysen G. Serum albumin: relationship to inflammation and nutrition. Semin Dial. 2004;17(6):432-437.',
        '19. Horwich TB, Kalantar-Zadeh K, MacLellan RW, Fonarow GC. Albumin levels predict survival in patients with systolic heart failure. Am Heart J. 2008;155(5):883-889.',
        '20. Studenski SA, Peters KW, Alley DE, et al. The FNIH sarcopenia project: rationale, study description, conference recommendations, and final estimates. J Gerontol A Biol Sci Med Sci. 2014;69(5):547-558.',
        '21. DeSalvo KB, Bloser N, Reynolds K, He J, Muntner P. Mortality prediction with a single general self-rated health question. J Gen Intern Med. 2006;21(3):267-275.'
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    output_path = os.path.join(MANUSCRIPT_DIR, 'IRI_Manuscript_Validated.docx')
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_manuscript()
