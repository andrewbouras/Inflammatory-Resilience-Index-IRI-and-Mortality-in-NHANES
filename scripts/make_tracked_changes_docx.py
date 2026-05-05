#!/usr/bin/env python3
"""
Build a Word docx redline diff between two markdown files.

Uses redlines.Redlines to compute word-level opcodes and python-docx to
write runs with strikethrough/red for deletions and bold/underlined/green
for insertions. The output is a real Word docx where Word and Google Docs
render the styling correctly (no fragile inline HTML span hacks).

Inputs:
  - /tmp/brahiri.md  (extracted via:
      pandoc IRI/brahiri.docx --from docx --to gfm --wrap=none -o /tmp/brahiri.md)
  - manuscript/IRI_Manuscript_Final.md  (the current canonical manuscript)

Output:
  - manuscript/IRI_Manuscript_TrackedChanges.docx

Dependencies (install once):
  python3 -m pip install --user redlines python-docx lxml
"""
from pathlib import Path
import re
from redlines import Redlines
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = Path("/Volumes/EXTERNAL SSD/localDocs/IRI")
BEFORE = Path("/tmp/brahiri.md")
AFTER = PROJECT / "manuscript" / "IRI_Manuscript_Final.md"
OUT_DOCX = PROJECT / "manuscript" / "IRI_Manuscript_TrackedChanges.docx"

RED = RGBColor(0xC0, 0x10, 0x10)      # deletion
GREEN = RGBColor(0x0A, 0x7D, 0x12)    # insertion
GRAY = RGBColor(0x55, 0x55, 0x55)

before_text = BEFORE.read_text()
after_text = AFTER.read_text()

diff = Redlines(before_text, after_text)
src_tokens = diff.source       # list of tokens for the original
dst_tokens = diff.test         # list of tokens for the new
opcodes = diff.opcodes          # [(tag, i1, i2, j1, j2), ...]

doc = Document()

# Page setup: letter, 0.75in margins
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = section.bottom_margin = Inches(0.75)
section.left_margin = section.right_margin = Inches(0.75)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Title block
title = doc.add_paragraph()
run = title.add_run("IRI Manuscript — Tracked Changes")
run.bold = True
run.font.size = Pt(16)

intro = doc.add_paragraph()
intro.add_run("Comparison: starting docx (").italic = True
code = intro.add_run("IRI/brahiri.docx")
code.font.name = "Consolas"; code.italic = True
intro.add_run(", the version PI Maciej Banach reviewed with comments) → current canonical manuscript (").italic = True
code = intro.add_run("manuscript/IRI_Manuscript_Final.md")
code.font.name = "Consolas"; code.italic = True
intro.add_run(", commit 89b8c07).").italic = True

# Reading guide
guide = doc.add_paragraph()
b = guide.add_run("Reading guide. ")
b.bold = True
r = guide.add_run("Strikethrough red text")
r.font.color.rgb = RED; r.font.strike = True
guide.add_run(" = removed since the version Banach reviewed. ")
r = guide.add_run("Underlined green text")
r.font.color.rgb = GREEN; r.bold = True; r.underline = True
guide.add_run(
    " = added since then. Unchanged text appears normally. "
    "Tables, figures, and structural elements may be flattened relative to the final docx; "
    "refer to manuscript/IRI_Manuscript_Final.docx for the authoritative typeset version."
)
doc.add_paragraph("─" * 60)


def add_run(p, text, kind):
    """Append a run to paragraph p with styling for kind in {'equal','insert','delete'}."""
    if not text:
        return
    r = p.add_run(text)
    if kind == "delete":
        r.font.color.rgb = RED
        r.font.strike = True
    elif kind == "insert":
        r.font.color.rgb = GREEN
        r.bold = True
        r.underline = True


# We have whitespace-split tokens that preserve newlines as standalone tokens.
# Build paragraphs by accumulating runs and breaking on token == "\n\n" or
# multiple consecutive newlines.

def emit_tokens(p_holder, tokens, kind):
    """
    Emit tokens (already stringified) into the document, handling paragraph
    breaks. p_holder is a 1-element list holding the current paragraph.
    Returns the (possibly new) current paragraph.
    """
    # Reassemble token text. redlines tokenizer keeps spaces as part of tokens
    # for some implementations; we'll just join verbatim and split on newlines.
    text = "".join(tokens)
    # Normalize Windows line endings
    text = text.replace("\r\n", "\n")
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            # paragraph break
            p_holder[0] = doc.add_paragraph()
        if part:
            add_run(p_holder[0], part, kind)


cur_paragraph = [doc.add_paragraph()]

for op in opcodes:
    tag, i1, i2, j1, j2 = op
    if tag == "equal":
        emit_tokens(cur_paragraph, src_tokens[i1:i2], "equal")
    elif tag == "delete":
        emit_tokens(cur_paragraph, src_tokens[i1:i2], "delete")
    elif tag == "insert":
        emit_tokens(cur_paragraph, dst_tokens[j1:j2], "insert")
    elif tag == "replace":
        # Treat as delete + insert for clarity
        emit_tokens(cur_paragraph, src_tokens[i1:i2], "delete")
        emit_tokens(cur_paragraph, dst_tokens[j1:j2], "insert")

# Add a footer separator
doc.add_paragraph("─" * 60)
foot = doc.add_paragraph()
fr = foot.add_run(
    "End of tracked-changes document. Generated automatically from a word-level diff "
    "between IRI/brahiri.docx (extracted to markdown via pandoc) and the current "
    "manuscript/IRI_Manuscript_Final.md."
)
fr.italic = True
fr.font.color.rgb = GRAY
fr.font.size = Pt(9)

doc.save(str(OUT_DOCX))
size = OUT_DOCX.stat().st_size
print(f"Wrote {OUT_DOCX} ({size:,} bytes)")
print(f"Source tokens: {len(src_tokens):,}")
print(f"Target tokens: {len(dst_tokens):,}")
print(f"Opcodes: {len(opcodes)}")
inserts = sum(1 for op in opcodes if op[0] in ("insert", "replace"))
deletes = sum(1 for op in opcodes if op[0] in ("delete", "replace"))
equals = sum(1 for op in opcodes if op[0] == "equal")
print(f"  equal:   {equals}")
print(f"  insert:  {sum(1 for op in opcodes if op[0]=='insert')}")
print(f"  delete:  {sum(1 for op in opcodes if op[0]=='delete')}")
print(f"  replace: {sum(1 for op in opcodes if op[0]=='replace')}")
