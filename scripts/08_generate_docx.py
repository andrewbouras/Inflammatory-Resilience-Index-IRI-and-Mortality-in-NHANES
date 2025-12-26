#!/usr/bin/env python3
"""
Generate IRI Manuscript as Word Document
Professional peer-reviewed format with proper paragraphs
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Paths
BASE_DIR = "/Users/andrewbouras/Documents/VishrutNHANES/Inflammatory Resilience Index (IRI) and Mortality in NHANES"
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
MANUSCRIPT_DIR = os.path.join(BASE_DIR, "manuscript")

def create_manuscript():
    """Create the complete manuscript document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph spacing
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_paragraph()
    title_run = title.add_run('Inflammatory Resilience Index Combining High-Sensitivity C-Reactive Protein, Albumin, and Appendicular Lean Mass Predicts Functional Limitations and Self-Rated Health in U.S. Adults: NHANES 2015-2020')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    authors = doc.add_paragraph('[Author names, degrees, and affiliations to be added]')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Correspondence
    corr = doc.add_paragraph()
    corr.add_run('Correspondence: ').bold = True
    corr.add_run('[Corresponding author name, address, email, phone]')
    
    doc.add_paragraph()
    
    # Word count
    wc = doc.add_paragraph()
    wc.add_run('Word count: ').bold = True
    wc.add_run('Abstract: 298; Manuscript: approximately 3,500')
    
    doc.add_page_break()
    
    # =========================================================================
    # ABSTRACT
    # =========================================================================
    abstract_title = doc.add_paragraph()
    abstract_title.add_run('ABSTRACT').bold = True
    abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Background
    p = doc.add_paragraph()
    p.add_run('Background: ').bold = True
    p.add_run('Inflammation, nutritional status, and muscle mass independently predict adverse health outcomes. An integrated index combining these domains may better capture overall physiologic resilience than individual biomarkers.')
    
    doc.add_paragraph()
    
    # Objectives
    p = doc.add_paragraph()
    p.add_run('Objectives: ').bold = True
    p.add_run('To develop an Inflammatory Resilience Index (IRI) combining high-sensitivity C-reactive protein (hs-CRP), serum albumin, and appendicular lean mass index (ALMI), and evaluate its association with functional outcomes in U.S. adults.')
    
    doc.add_paragraph()
    
    # Methods
    p = doc.add_paragraph()
    p.add_run('Methods: ').bold = True
    p.add_run('Using NHANES 2015-2020 with dual-energy X-ray absorptiometry (DEXA) data, we constructed the IRI as the sum of inverted standardized hs-CRP, standardized albumin, and sex-specific standardized ALMI, where higher values indicate better resilience. Among 2,729 adults aged 20 years or older with complete data, we examined associations between IRI quartiles and functional outcomes, including fair or poor self-rated health, difficulty walking one-quarter mile, and depression defined as Patient Health Questionnaire-9 score of 10 or higher, using survey-weighted logistic regression adjusted for age, sex, and race/ethnicity.')
    
    doc.add_paragraph()
    
    # Results
    p = doc.add_paragraph()
    p.add_run('Results: ').bold = True
    p.add_run('The IRI ranged from -5.2 to 5.2 with a mean of 0.74. Participants in the lowest IRI quartile (Q1) were older, more often female, had higher hs-CRP, lower albumin, and lower ALMI z-scores compared with Q4. The prevalence of fair or poor health decreased from 19.2% in Q1 to 9.2% in Q4, and difficulty walking decreased from 12.4% to 2.5%. In adjusted models, each 1-unit IRI increase was associated with 19% lower odds of fair or poor health (odds ratio 0.81; 95% confidence interval 0.74-0.89; P<0.001) and 22% lower odds of walking difficulty (odds ratio 0.78; 95% confidence interval 0.67-0.91; P=0.005). Comparing Q1 to Q4, odds were 2.07-fold higher for fair or poor health (95% confidence interval 1.42-3.01; P=0.004) and 4.51-fold higher for walking difficulty (95% confidence interval 1.96-10.42; P=0.006).')
    
    doc.add_paragraph()
    
    # Conclusions
    p = doc.add_paragraph()
    p.add_run('Conclusions: ').bold = True
    p.add_run('A novel Inflammatory Resilience Index integrating inflammation, nutritional reserve, and muscle mass is significantly associated with self-rated health and mobility limitations in U.S. adults. Lower IRI identifies individuals with substantially higher odds of poor functional outcomes, supporting its potential utility as a composite marker of physiologic reserve.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Keywords: ').italic = True
    p.add_run('Inflammation; C-reactive protein; albumin; sarcopenia; functional status; NHANES')
    
    doc.add_page_break()
    
    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    intro_title = doc.add_paragraph()
    intro_title.add_run('INTRODUCTION').bold = True
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('Systemic inflammation, nutritional status, and skeletal muscle mass represent three interconnected physiologic domains that independently predict adverse health outcomes including mortality, cardiovascular disease, and functional decline.1-3 High-sensitivity C-reactive protein (hs-CRP), a marker of low-grade systemic inflammation, is associated with increased cardiovascular risk even after adjustment for traditional risk factors.4 Serum albumin, reflecting both hepatic synthetic function and nutritional adequacy, predicts mortality across diverse populations.5 Appendicular lean mass, a measure of skeletal muscle, declines with aging and is a key component of sarcopenia and frailty syndromes.6')
    
    doc.add_paragraph()
    
    doc.add_paragraph('While each biomarker provides prognostic information independently, the interplay between inflammation, nutrition, and muscle mass suggests that integrated assessment may better capture overall physiologic resilience. Chronic inflammation promotes catabolism and muscle wasting, malnutrition impairs immune function and muscle protein synthesis, and sarcopenia is associated with increased inflammatory markers.7-9 This bidirectional relationship creates a vicious cycle that accelerates functional decline in vulnerable individuals.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Previous composite indices have combined subsets of these domains. The Glasgow Prognostic Score uses CRP and albumin to predict cancer outcomes.10 Frailty indices incorporate measures of strength, nutrition, and functional status.11 However, a simple composite specifically integrating inflammation, nutrition, and objective muscle mass measurement has not been widely evaluated in population-based samples.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('We hypothesized that an Inflammatory Resilience Index (IRI) combining hs-CRP, serum albumin, and appendicular lean mass index (ALMI) from dual-energy X-ray absorptiometry (DEXA) would identify individuals with poor self-rated health and functional limitations. Using nationally representative data from NHANES 2015-2020, we aimed to construct the IRI as a simple additive z-score composite, characterize the demographic and clinical profiles of adults across IRI quartiles, and evaluate associations between IRI and functional outcomes including self-rated health, mobility limitations, and depressive symptoms.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('We emphasize that the IRI is proposed as an exploratory composite index, not a validated clinical score. This analysis represents an initial step toward understanding whether integrated assessment of these three domains provides incremental information beyond individual biomarkers.')
    
    doc.add_page_break()
    
    # =========================================================================
    # METHODS
    # =========================================================================
    methods_title = doc.add_paragraph()
    methods_title.add_run('METHODS').bold = True
    methods_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Data Source
    p = doc.add_paragraph()
    p.add_run('Data Source and Study Population').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('This analysis used data from the National Health and Nutrition Examination Survey (NHANES), a nationally representative survey of the civilian, non-institutionalized U.S. population conducted by the National Center for Health Statistics.12 NHANES employs a complex, multistage probability sampling design with oversampling of certain subgroups to ensure reliable estimates. Participants undergo standardized interviews, physical examinations, and laboratory assessments.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('We combined data from NHANES cycles 2015-2016 and 2017-2020 (pre-pandemic), which include DEXA body composition measurements. The analytic sample included adults aged 20 years or older with complete data for hs-CRP at 10 mg/L or less (to exclude acute infection), serum albumin, and DEXA-derived appendicular lean mass. Participants with hs-CRP greater than 10 mg/L were excluded to focus on chronic low-grade inflammation rather than acute infectious or inflammatory conditions.13 The final eligible cohort included 2,729 participants with complete IRI components and functional outcome data.')
    
    doc.add_paragraph()
    
    # IRI Construction
    p = doc.add_paragraph()
    p.add_run('Inflammatory Resilience Index Construction').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('The IRI was constructed as an additive composite of three standardized components: inverted standardized log-transformed hs-CRP, standardized serum albumin, and sex-specific standardized ALMI. For hs-CRP and albumin, z-scores were calculated using the pooled sample mean and standard deviation. For ALMI, sex-specific standardization was used to account for known differences in muscle mass between men and women. The hs-CRP component was inverted so that lower inflammation yields higher scores. Higher IRI values indicate better inflammatory resilience, reflecting lower inflammation, higher albumin, and greater muscle mass. This formulation assigns equal weight to each domain and is designed for simplicity and interpretability. The composite IRI was not re-centered after construction; thus, the sample mean deviates slightly from zero due to eligibility restrictions and missing data patterns.')
    
    doc.add_paragraph()
    
    # Component Measurements
    p = doc.add_paragraph()
    p.add_run('Component Measurements').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('High-sensitivity C-reactive protein was measured in serum using latex-enhanced nephelometry at a central laboratory. Values were reported in mg/L. Serum albumin was measured using the bromocresol purple method in the standard NHANES biochemistry panel and reported in g/dL. Appendicular lean mass index was calculated from whole-body DEXA scans using a Hologic Discovery model A densitometer. Appendicular lean mass was defined as the sum of lean soft tissue mass in the arms and legs, excluding bone mineral content. ALMI was calculated as appendicular lean mass in kilograms divided by height squared in meters. For the IRI, ALMI was converted to sex-specific z-scores to account for known sex differences in muscle mass while facilitating combination with other index components.')
    
    doc.add_paragraph()
    
    # Functional Outcomes
    p = doc.add_paragraph()
    p.add_run('Functional Outcomes').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Primary outcomes were obtained from NHANES questionnaire data. For self-rated health, participants rated their general health as excellent, very good, good, fair, or poor, and we created a binary outcome for fair or poor health versus excellent, very good, or good. For mobility limitation, participants reported difficulty walking a quarter mile, with those reporting some difficulty, much difficulty, or unable to do classified as having walking difficulty. For depression, the Patient Health Questionnaire-9 (PHQ-9) was administered, and a score of 10 or higher was used to identify moderate-to-severe depressive symptoms, consistent with standard clinical cutoffs.14')
    
    doc.add_paragraph()
    
    # Statistical Analysis
    p = doc.add_paragraph()
    p.add_run('Statistical Analysis').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('All analyses incorporated NHANES survey weights, strata, and primary sampling units to account for the complex survey design. Examination weights were scaled for combined cycles per NCHS guidelines.15 Because DEXA is performed in a subsample of examined participants, estimates are nationally representative of U.S. adults eligible for and completing DEXA body composition assessment; older adults and those with mobility limitations may be underrepresented.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Baseline characteristics were compared across IRI quartiles using survey-weighted means and proportions. Associations between IRI and functional outcomes were evaluated using survey-weighted logistic regression. We used three modeling approaches: continuous IRI examining the odds ratio per 1-unit increase, quartile analysis examining odds ratios for Q1, Q2, and Q3 compared with Q4 as the reference representing highest resilience, and sensitivity analysis with fully adjusted models including body mass index, diabetes, hypertension, and smoking. Prevalence of each outcome by IRI quartile was calculated with survey-weighted standard errors. A two-sided P value less than 0.05 was considered statistically significant. Analyses were performed using R version 4.3 with the survey package.16')
    
    doc.add_page_break()
    
    # =========================================================================
    # RESULTS
    # =========================================================================
    results_title = doc.add_paragraph()
    results_title.add_run('RESULTS').bold = True
    results_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Study Population
    p = doc.add_paragraph()
    p.add_run('Study Population').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Of 25,531 NHANES 2015-2020 participants, 2,729 adults aged 20 years or older met eligibility criteria with complete IRI components. The mean age was 39.6 years (standard deviation 11.4), 48.6% were female, and the racial and ethnic distribution was 61.5% non-Hispanic White, 11.0% non-Hispanic Black, 8.4% Mexican American, 10.8% Asian, and 8.3% other. The IRI ranged from -5.16 to 5.24, with a mean of 0.74 and standard deviation of 1.48. Quartile boundaries were Q1 at -0.28 or less representing lowest resilience, Q2 from -0.28 to 0.83, Q3 from 0.83 to 1.90, and Q4 greater than 1.90 representing highest resilience.')
    
    doc.add_paragraph()
    
    # Characteristics by Quartile
    p = doc.add_paragraph()
    p.add_run('Characteristics by IRI Quartile').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Participants in the lowest IRI quartile differed substantially from those in the highest quartile (Table 1). Q1 participants were older with a mean age of 42.3 versus 34.3 years (P<0.001), more often female at 68.6% versus 25.1% (P<0.001), and had higher body mass index at 29.5 versus 27.4 kg/m2 (P<0.001). IRI component profiles showed marked differences: hs-CRP was 5.3-fold higher in Q1 at 4.08 versus 0.77 mg/L, albumin was lower at 4.12 versus 4.70 g/dL, and ALMI z-scores were lower at -0.40 versus 0.53 (all P<0.001). Diabetes was more prevalent in Q1 at 15.6% versus 5.7% (P<0.001), while hypertension prevalence was similar across quartiles.')
    
    doc.add_paragraph()
    
    # Functional Outcomes
    p = doc.add_paragraph()
    p.add_run('Functional Outcomes by IRI Quartile').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Clear gradients in functional outcome prevalence were observed across IRI quartiles (Table 2). The prevalence of fair or poor self-rated health decreased from 19.2% in Q1 to 17.5% in Q2, 13.8% in Q3, and 9.2% in Q4, representing an absolute difference of 10.0 percentage points between extreme quartiles. Difficulty walking one-quarter mile decreased from 12.4% in Q1 to 8.7% in Q2, 7.8% in Q3, and 2.5% in Q4, an absolute difference of 9.9 percentage points. Depression prevalence showed a more modest gradient, decreasing from 8.5% in Q1 to 5.3% in Q4.')
    
    doc.add_paragraph()
    
    # Multivariable Associations
    p = doc.add_paragraph()
    p.add_run('Multivariable Associations').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('In continuous IRI analysis, each 1-unit increase in IRI was associated with significantly lower odds of adverse functional outcomes after adjustment for age, sex, and race/ethnicity (Table 2). For fair or poor health, the odds ratio was 0.81 (95% confidence interval 0.74-0.89; P<0.001). For walking difficulty, the odds ratio was 0.78 (95% confidence interval 0.67-0.91; P=0.005). For depression, the odds ratio was 0.90 (95% confidence interval 0.77-1.05; P=0.14), which was not statistically significant.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('In quartile analysis with Q4 as the reference, participants in the lowest quartile had significantly elevated odds of poor outcomes. For fair or poor health, Q1 versus Q4 yielded an odds ratio of 2.07 (95% confidence interval 1.42-3.01; P=0.004). For walking difficulty, Q1 versus Q4 yielded an odds ratio of 4.51 (95% confidence interval 1.96-10.42; P=0.006). For depression, Q1 versus Q4 yielded an odds ratio of 1.40 (95% confidence interval 0.79-2.49; P=0.19), which was not statistically significant. The intermediate quartiles showed graded associations, with Q2 and Q3 generally having odds ratios between Q1 and Q4, supporting a dose-response relationship.')
    
    doc.add_paragraph()
    
    # Sensitivity Analyses
    p = doc.add_paragraph()
    p.add_run('Sensitivity Analyses').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('In fully adjusted models including body mass index, diabetes, hypertension, and smoking, associations were modestly attenuated but remained significant for self-rated health, with continuous IRI yielding an odds ratio of 0.85 (95% confidence interval 0.76-0.95; P=0.02). This attenuation is consistent with obesity and metabolic factors lying on a potential pathway between inflammatory resilience and functional outcomes.')
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE 1
    # =========================================================================
    table1_title = doc.add_paragraph()
    table1_title.add_run('Table 1. Baseline Characteristics by IRI Quartile').bold = True
    
    doc.add_paragraph()
    
    table1 = doc.add_table(rows=13, cols=6)
    table1.style = 'Table Grid'
    
    headers = ['Characteristic', 'Q1 (n=652)', 'Q2 (n=757)', 'Q3 (n=743)', 'Q4 (n=577)', 'P value']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['Age, years', '42.3 (11.2)', '41.4 (11.4)', '39.6 (10.9)', '34.3 (10.7)', '<0.001'],
        ['Female, %', '68.6', '57.5', '40.6', '25.1', '<0.001'],
        ['BMI, kg/m2', '29.5 (7.0)', '29.2 (6.2)', '29.2 (6.7)', '27.4 (6.3)', '<0.001'],
        ['hs-CRP, mg/L', '4.08 (2.62)', '2.62 (2.07)', '1.77 (1.75)', '0.77 (1.12)', '<0.001'],
        ['Albumin, g/dL', '4.12 (0.28)', '4.32 (0.24)', '4.47 (0.22)', '4.70 (0.23)', '<0.001'],
        ['ALMI z-score*', '-0.40 (1.09)', '0.22 (0.83)', '0.49 (0.77)', '0.53 (0.81)', '<0.001'],
        ['Diabetes, %', '15.6', '11.8', '9.6', '5.7', '<0.001'],
        ['Hypertension, %', '41.0', '40.3', '39.4', '33.1', '0.20'],
        ['CVD history, %', '5.4', '3.0', '3.8', '2.3', '0.12'],
        ['Current smoker, %', '25.8', '18.7', '21.9', '16.5', '0.03'],
        ['IRI, mean (SD)', '-1.20 (0.71)', '0.30 (0.32)', '1.34 (0.31)', '2.72 (0.69)', '--']
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table1.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Values are mean (SD) or percentage. *ALMI z-scores are sex-specific standardized values. Abbreviations: BMI, body mass index; CVD, cardiovascular disease; hs-CRP, high-sensitivity C-reactive protein; ALMI, appendicular lean mass index; IRI, Inflammatory Resilience Index.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE 2
    # =========================================================================
    table2_title = doc.add_paragraph()
    table2_title.add_run('Table 2. Association of Inflammatory Resilience Index with Functional Outcomes').bold = True
    
    doc.add_paragraph()
    
    # Panel A
    p = doc.add_paragraph()
    p.add_run('Panel A: Prevalence by IRI Quartile').italic = True
    
    doc.add_paragraph()
    
    table2a = doc.add_table(rows=4, cols=5)
    table2a.style = 'Table Grid'
    
    headers2a = ['Outcome', 'Q1', 'Q2', 'Q3', 'Q4']
    for i, h in enumerate(headers2a):
        table2a.rows[0].cells[i].text = h
        table2a.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    prev_data = [
        ['Fair/poor health, % (SE)', '19.2 (2.0)', '17.5 (1.6)', '13.8 (1.9)', '9.2 (1.0)'],
        ['Walking difficulty, % (SE)', '12.4 (1.7)', '8.7 (1.4)', '7.8 (1.3)', '2.5 (0.6)'],
        ['Depression, % (SE)', '8.5 (0.9)', '7.0 (1.3)', '6.6 (0.9)', '5.3 (0.6)']
    ]
    
    for i, row in enumerate(prev_data):
        for j, cell in enumerate(row):
            table2a.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    
    # Panel B
    p = doc.add_paragraph()
    p.add_run('Panel B: Adjusted Odds Ratios').italic = True
    
    doc.add_paragraph()
    
    table2b = doc.add_table(rows=4, cols=5)
    table2b.style = 'Table Grid'
    
    headers2b = ['Outcome', 'Per 1-unit IRI (95% CI)', 'P', 'Q1 vs Q4 (95% CI)', 'P']
    for i, h in enumerate(headers2b):
        table2b.rows[0].cells[i].text = h
        table2b.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    or_data = [
        ['Fair/poor health', '0.81 (0.74-0.89)', '<0.001', '2.07 (1.42-3.01)', '0.004'],
        ['Walking difficulty', '0.78 (0.67-0.91)', '0.005', '4.51 (1.96-10.42)', '0.006'],
        ['Depression', '0.90 (0.77-1.05)', '0.14', '1.40 (0.79-2.49)', '0.19']
    ]
    
    for i, row in enumerate(or_data):
        for j, cell in enumerate(row):
            table2b.rows[i+1].cells[j].text = cell
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Models adjusted for age, sex, and race/ethnicity. Abbreviations: CI, confidence interval; IRI, Inflammatory Resilience Index.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # FIGURES
    # =========================================================================
    fig_title = doc.add_paragraph()
    fig_title.add_run('FIGURES').bold = True
    fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Figure 1
    p = doc.add_paragraph()
    p.add_run('Figure 1. Forest Plot of IRI Quartile Associations with Functional Outcomes').bold = True
    
    doc.add_paragraph()
    
    fig1_path = os.path.join(OUTPUT_DIR, 'figure1_forest_plot.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Odds ratios comparing the lowest IRI quartile (Q1) with the highest (Q4, reference) for three functional outcomes, adjusted for age, sex, and race/ethnicity. Red points indicate statistically significant associations (P<0.05). Error bars represent 95% confidence intervals.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # Figure 2
    p = doc.add_paragraph()
    p.add_run('Figure 2. Prevalence of Functional Outcomes by IRI Quartile').bold = True
    
    doc.add_paragraph()
    
    fig2_path = os.path.join(OUTPUT_DIR, 'figure2_prevalence_by_quartile.png')
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Survey-weighted prevalence with standard error bars for fair or poor self-rated health, difficulty walking one-quarter mile, and depression across IRI quartiles. Q1 represents lowest inflammatory resilience and Q4 represents highest resilience.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # Figure 3
    p = doc.add_paragraph()
    p.add_run('Figure 3. IRI Component Profiles by Quartile').bold = True
    
    doc.add_paragraph()
    
    fig3_path = os.path.join(OUTPUT_DIR, 'figure3_iri_components.png')
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Mean values of high-sensitivity C-reactive protein (mg/L), serum albumin (g/dL), and ALMI z-score across IRI quartiles, illustrating the distribution of each component contributing to the composite index.').font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # DISCUSSION
    # =========================================================================
    disc_title = doc.add_paragraph()
    disc_title.add_run('DISCUSSION').bold = True
    disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('In this nationally representative sample of U.S. adults, we developed and evaluated an Inflammatory Resilience Index integrating hs-CRP, serum albumin, and DEXA-derived appendicular lean mass. Lower IRI was significantly associated with fair or poor self-rated health and mobility limitations after adjustment for age, sex, and race/ethnicity. Adults in the lowest IRI quartile had approximately 2-fold higher odds of poor self-rated health and 4.5-fold higher odds of walking difficulty compared with the highest quartile. These findings support the concept that integrated assessment across inflammation, nutrition, and muscle mass domains may identify individuals at risk for functional decline.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('The associations observed are biologically plausible given the well-established interrelationships between the IRI components. Chronic low-grade inflammation promotes muscle catabolism through activation of the ubiquitin-proteasome pathway and suppression of protein synthesis.7 Inflammatory cytokines including interleukin-6 and tumor necrosis factor-alpha directly impair myocyte function and promote sarcopenia.17 Albumin, beyond its role as a nutritional marker, functions as an antioxidant and anti-inflammatory mediator; hypoalbuminemia reflects both inadequate nutritional intake and increased catabolism during inflammatory states.18')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Prior NHANES analyses have examined individual IRI components in isolation. Elevated hs-CRP is associated with increased cardiovascular and all-cause mortality, even after adjustment for traditional risk factors.4 Low serum albumin predicts mortality across age groups and is incorporated into prognostic scores for heart failure and chronic kidney disease.19 Low appendicular lean mass identifies individuals at risk for falls, fractures, and functional decline.20 Our analysis extends this literature by demonstrating that a simple additive composite of standardized components is associated with self-reported functional outcomes.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('The strong association between IRI and self-rated health is particularly notable, as self-rated health is itself a powerful predictor of mortality and health care utilization.21 Identifying modifiable contributors to poor self-rated health could inform targeted interventions.')
    
    doc.add_paragraph()
    
    # Limitations
    p = doc.add_paragraph()
    p.add_run('Limitations').italic = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('Several limitations warrant consideration. First, the cross-sectional design precludes causal inference. We cannot determine whether low IRI causes poor functional outcomes, whether functional limitations lead to inflammation and muscle loss, or whether both reflect underlying comorbid conditions. Prospective validation with incident outcomes is needed.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Second, the analytic sample was restricted to participants with DEXA data, which represents a younger and healthier subset of NHANES with a mean age of 39.6 years. Results may not generalize to older or sicker populations.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Third, ALMI values were standardized to sex-specific z-scores rather than using established sarcopenia cutoffs. This approach facilitates combination with other components but limits comparison with prior sarcopenia literature.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Fourth, functional outcomes were self-reported and subject to reporting bias. Objective measures of physical function such as gait speed or grip strength would strengthen causal inference.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Fifth, the IRI is proposed as an exploratory composite. Its components are weighted equally, and the optimal weighting scheme was not empirically derived. The index has not been validated in external cohorts. An exploratory mortality analysis is presented in the Supplemental Materials; however, only 20 deaths occurred, and mortality findings are severely underpowered and should be interpreted with caution.')
    
    doc.add_page_break()
    
    # =========================================================================
    # CONCLUSIONS
    # =========================================================================
    conc_title = doc.add_paragraph()
    conc_title.add_run('CONCLUSIONS').bold = True
    conc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_paragraph('In this nationally representative sample of U.S. adults, the Inflammatory Resilience Index, a simple composite of hs-CRP, serum albumin, and appendicular lean mass, was significantly associated with self-rated health and mobility limitations. Adults in the lowest IRI quartile had 2-fold higher odds of fair or poor health and 4.5-fold higher odds of walking difficulty compared with the highest quartile. These findings support IRI as an integrative marker of physiologic resilience capturing inflammation, nutrition, and muscle mass domains.')
    
    doc.add_paragraph()
    
    doc.add_paragraph('The IRI is proposed as an exploratory index warranting validation in prospective cohorts with hard clinical endpoints. If confirmed, the IRI could inform risk stratification and identify targets for intervention in individuals at risk for functional decline.')
    
    doc.add_page_break()
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    ref_title = doc.add_paragraph()
    ref_title.add_run('REFERENCES').bold = True
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    references = [
        '1. Libby P, Ridker PM, Hansson GK. Inflammation in atherosclerosis: from pathophysiology to practice. J Am Coll Cardiol. 2009;54(23):2129-2138.',
        '2. Ridker PM, Cushman M, Stampfer MJ, et al. Inflammation, aspirin, and the risk of cardiovascular disease in apparently healthy men. N Engl J Med. 1997;336(14):973-979.',
        '3. Cruz-Jentoft AJ, Bahat G, Bauer J, et al. Sarcopenia: revised European consensus on definition and diagnosis. Age Ageing. 2019;48(1):16-31.',
        '4. Ridker PM. High-sensitivity C-reactive protein: potential adjunct for global risk assessment in the primary prevention of cardiovascular disease. Circulation. 2001;103(13):1813-1818.',
        '5. Goldwasser P, Feldman J. Association of serum albumin and mortality risk. J Clin Epidemiol. 1997;50(6):693-703.',
        '6. Bauer J, Morley JE, Schols AMWJ, et al. Sarcopenia: a time for action. J Cachexia Sarcopenia Muscle. 2019;10(5):956-961.',
        '7. Schaap LA, Pluijm SMF, Deeg DJH, Visser M. Inflammatory markers and loss of muscle mass (sarcopenia) and strength. Am J Med. 2006;119(6):526.e9-17.',
        '8. Dalle S, Rossmeislova L, Bhide M. The role of inflammation in age-related sarcopenia. Front Physiol. 2017;8:1045.',
        '9. Wilson D, Jackson T, Sapey E, Lord JM. Frailty and sarcopenia: the potential role of an aged immune system. Ageing Res Rev. 2017;36:1-10.',
        '10. McMillan DC. The systemic inflammation-based Glasgow Prognostic Score: a decade of experience in patients with cancer. Cancer Treat Rev. 2013;39(5):534-540.',
        '11. Rockwood K, Song X, MacKnight C, et al. A global clinical measure of fitness and frailty in elderly people. CMAJ. 2005;173(5):489-495.',
        '12. Centers for Disease Control and Prevention (CDC). National Health and Nutrition Examination Survey. https://www.cdc.gov/nchs/nhanes/index.htm. Accessed December 2025.',
        '13. Pearson TA, Mensah GA, Alexander RW, et al. Markers of inflammation and cardiovascular disease: application to clinical and public health practice. Circulation. 2003;107(3):499-511.',
        '14. Kroenke K, Spitzer RL, Williams JB. The PHQ-9: validity of a brief depression severity measure. J Gen Intern Med. 2001;16(9):606-613.',
        '15. Johnson CL, Paulose-Ram R, Ogden CL, et al. National Health and Nutrition Examination Survey: analytic guidelines, 1999-2010. Vital Health Stat 2. 2013;(161):1-24.',
        '16. Lumley T. Analysis of complex survey samples. J Stat Softw. 2004;9(8):1-19.',
        '17. Schaap LA, Pluijm SMF, Deeg DJH, et al. Higher inflammatory marker levels in older persons: associations with 5-year change in muscle mass and muscle strength. J Gerontol A Biol Sci Med Sci. 2009;64(11):1183-1189.',
        '18. Don BR, Kaysen G. Serum albumin: relationship to inflammation and nutrition. Semin Dial. 2004;17(6):432-437.',
        '19. Horwich TB, Kalantar-Zadeh K, MacLellan RW, Fonarow GC. Albumin levels predict survival in patients with systolic heart failure. Am Heart J. 2008;155(5):883-889.',
        '20. Studenski SA, Peters KW, Alley DE, et al. The FNIH sarcopenia project: rationale, study description, conference recommendations, and final estimates. J Gerontol A Biol Sci Med Sci. 2014;69(5):547-558.',
        '21. DeSalvo KB, Bloser N, Reynolds K, He J, Muntner P. Mortality prediction with a single general self-rated health question. J Gen Intern Med. 2006;21(3):267-275.'
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    output_path = os.path.join(MANUSCRIPT_DIR, 'IRI_Manuscript_Final.docx')
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_manuscript()
